"""
Report exporter.

Renders a MonthlyReport to portable formats. Markdown is always available (dependency-free);
this is what the dashboard's "export" / copy action uses. A docx export hook is provided and
degrades gracefully if python-docx isn't available.
"""

from __future__ import annotations

import io


def to_markdown(report) -> str:
    lines = [f"# iTrix Monthly Report — {report.month}", ""]
    for section in report.sections or []:
        lines.append(f"## {section.get('title', '')}")
        lines.append(section.get("body", ""))
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def to_docx_bytes(report) -> bytes | None:
    """Return a .docx as bytes, or None if python-docx is unavailable."""
    try:
        import docx  # noqa: PLC0415
    except Exception:  # noqa: BLE001
        return None
    document = docx.Document()
    document.add_heading(f"iTrix Monthly Report — {report.month}", level=0)
    for section in report.sections or []:
        document.add_heading(section.get("title", ""), level=1)
        document.add_paragraph(section.get("body", ""))
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

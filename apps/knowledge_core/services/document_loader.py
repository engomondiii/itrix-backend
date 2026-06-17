"""
Document loader.

Reads a registered document's raw text from a local path or an ``s3://`` URI. Supports
``.docx``, ``.pdf``, ``.txt`` and ``.md``. The heavy parsers (python-docx, pypdf) are
imported lazily so the module loads even if they're absent; a clear error is raised only
if a format that needs them is actually loaded.
"""

from __future__ import annotations

import io
import logging
import os

from storage.backends import read_file_bytes

logger = logging.getLogger("itrix")

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".txt", ".md", ".markdown"}


class UnsupportedDocument(ValueError):
    pass


def _ext(path: str) -> str:
    return os.path.splitext(path)[1].lower()


def _load_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _load_docx(data: bytes) -> str:
    try:
        import docx  # python-docx  # noqa: PLC0415
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedDocument(
            "python-docx is required to read .docx files."
        ) from exc
    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for para in document.paragraphs:
        text = (para.text or "").strip()
        style = (para.style.name if para.style else "") or ""
        if not text:
            continue
        # Preserve heading markers so the chunker can split on them.
        if style.lower().startswith("heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "1"
            parts.append(f"{'#' * min(int(level), 6)} {text}")
        else:
            parts.append(text)
    # Include simple table text too.
    for table in getattr(document, "tables", []):
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _load_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # noqa: PLC0415
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedDocument("pypdf is required to read .pdf files.") from exc
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            continue
    return "\n\n".join(p.strip() for p in pages if p.strip())


def load_document_text(source_ref: str) -> str:
    """Return the plain text of the document at ``source_ref`` (local path or s3://)."""
    ext = _ext(source_ref)
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocument(f"Unsupported document type: {ext or '(none)'}")

    data = read_file_bytes(source_ref)

    if ext in {".txt", ".md", ".markdown"}:
        text = _load_txt(data)
    elif ext == ".docx":
        text = _load_docx(data)
    elif ext == ".pdf":
        text = _load_pdf(data)
    else:  # pragma: no cover - guarded above
        raise UnsupportedDocument(f"Unsupported document type: {ext}")

    logger.info("Loaded %s (%d chars) from %s", ext, len(text), source_ref)
    return text

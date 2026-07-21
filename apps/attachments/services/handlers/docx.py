"""DOCX text extraction, including table cells."""

from __future__ import annotations

from io import BytesIO

from apps.attachments.services.handlers import ExtractionResult, metadata_only


def extract(data: bytes, *, filename: str = "", limit: int = 400_000) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError:
        return metadata_only("docx", "docx reader unavailable")

    try:
        document = Document(BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        # Tables carry the actual numbers in most technical documents, so a paragraph-only
        # extraction would miss the part a visitor most wants us to read.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        if not text:
            return metadata_only("docx", "no text content")
        return ExtractionResult(text=text[:limit], handler="docx", truncated=len(text) > limit)
    except Exception as exc:  # noqa: BLE001
        return metadata_only("docx", f"unreadable: {exc}"[:200])
